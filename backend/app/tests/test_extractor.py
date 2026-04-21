import os
import tempfile
import pytest
from app.services.extractor import extract_text, _normalize


def _write(content, suffix):
    fd, path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    return path


def test_txt_extraction():
    path = _write("Hello  world\n\n\n\nTest", ".txt")
    text = extract_text(path, ".txt")
    os.unlink(path)
    assert "Hello world" in text
    assert "Test" in text


def test_html_extraction():
    path = _write("<html><body><p>Hello <b>world</b></p><script>var x=1;</script></body></html>", ".html")
    text = extract_text(path, ".html")
    os.unlink(path)
    assert "Hello" in text
    assert "world" in text
    assert "var x" not in text


def test_md_extraction():
    path = _write("# Title\n\nSome **bold** text.", ".md")
    text = extract_text(path, ".md")
    os.unlink(path)
    assert "Title" in text
    assert "bold" in text


def test_normalize_whitespace():
    assert _normalize("  hello   world  ") == "hello world"


def test_normalize_blank_lines():
    result = _normalize("a\n\n\n\n\nb")
    assert "\n\n\n" not in result


def test_docx_extraction(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    path = str(tmp_path / "test.docx")
    doc.save(path)
    text = extract_text(path, ".docx")
    assert "First paragraph" in text
    assert "Second paragraph" in text
